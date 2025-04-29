from calendar import month
import logging
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import ElementClickInterceptedException, NoSuchElementException
from selenium.webdriver.chrome.options import Options
from persian_tools import digits
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from persiantools.jdatetime import JalaliDate
import os, json, time
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import urllib.request
from datetime import datetime
from Eitaa_api import send_to_Eitaa
from phones_urls import digi_urls, techno_urls
import traceback
import argparse

# Suppressing unnecessary error messages
os.environ["GRPC_VERBOSITY"] = "ERROR"
os.environ["GLOG_minloglevel"] = "2"
logger = logging.getLogger('selenium')
t_prices = []
d_prices = []

GREEN = "\033[32m"
RED = "\033[31m"
YELLOW = "\033[33m"
RESET = "\033[0m"

Test_Mode = False

def driver_setup():
    chrome_options = Options()
    chrome_options.add_argument('--headless')
    # chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--disable-popup-blocking")
    chrome_options.add_argument("--start-maximized")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-extensions")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-infobars")
    chrome_options.add_argument("--disable-browser-side-navigation")
    chrome_options.add_argument("--disable-images")
    chrome_options.add_argument("--log-level=3")
    # chrome_options.add_argument("user-data-dir=./cache")
    driver = webdriver.Chrome(options=chrome_options)

    return driver


xpath_for_black_techno = [
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[1]/div/p[contains(text() , "مشکی")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[2]/div/p[contains(text() , "مشکی")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[3]/div/p[contains(text() , "مشکی")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[4]/div/p[contains(text() , "مشکی")]'
]
xpath_for_darkblue = [
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[1]/div/p[contains(text() , "سرمه‌ای")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[2]/div/p[contains(text() , "سرمه‌ای")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[3]/div/p[contains(text() , "سرمه‌ای")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[4]/div/p[contains(text() , "سرمه‌ای")]'
]

xpath_for_white = [
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[1]/div/p[contains(text() , "سفید")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[2]/div/p[contains(text() , "سفید")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[3]/div/p[contains(text() , "سفید")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[4]/div/p[contains(text() , "سفید")]'
]


xpath_for_price_techno = {
    '1': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div/div[3]/div[3]/div[2]/div/div/p',
    '2': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div/div[3]/div[5]/div/div/div/p',
    '3': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div[2]/div[3]/div[3]/div/div/div/p[2]',
    '4': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div/div[3]/div[3]/div/div/div/p',
    '5': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div/div[3]/div[4]/div/div/div/p',
    '6': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div[2]/div[3]/div[2]/div[2]/div/div/p[2]',
    '7': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div/div[3]/div[2]/div[2]/div/div/p',
    '8': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div/div[3]/div[2]/div/div/div/p',
    '9': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div[2]/div[3]/div[2]/div/div/div/p[2]',
    '10': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div[2]/div[3]/div[4]/div/div/div/p[2]'
}

def check_internet_connection():
    try: # testing the connection by pinging google
        urllib.request.urlopen('https://www.google.com/', timeout=5)
        return True
    except Exception as e:
        print(f"{YELLOW}[!]{RESET}Connection problem --> pinging google had an unknown exception, code error : {e}")
        return False

def wait_for_connection(max_retries=10, retry_delay=10):
    # waiting for user to reconnect the connection
    retries = 0
    while retries < max_retries:
        if check_internet_connection():
            if retries > 0:
                print(f"{GREEN}[✓]{RESET}Internet is connected....")
            return True
        else:
            retries += 1
            print(f"{YELLOW}[!]{RESET}No internet connection. Retrying in {retry_delay} seconds... ({retries}/{max_retries})")
            time.sleep(retry_delay)

    print("Failed to reconnect after multiple attempts.")
    return False


if len(digi_urls) != len(techno_urls):
    raise Exception(f"{RED}[!]{RESET}The number of urls for technolife and digikala are different")
else:
    urls_len = len(digi_urls)
    phone_models = []
    for key in digi_urls.keys():
        phone_models.append(key)

def deny(btn, driver):
    try:
        # Wait for the 'deny' button to appear
        deny__btn = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.ID, "webpush-onsite"))
        )
        iframe = driver.find_element(By.ID, 'webpush-onsite')
        driver.switch_to.frame(iframe)
        
        # Try clicking the deny button
        try:
            # click(driver, (By.ID, "webpush-onsite"))
            deny__btn.click()
        except Exception as e:
            if isinstance(e, ElementClickInterceptedException):
                logger.debug("Debug : ElementClickInterceptedException")
            else:
                logger.debug(f"Debug : Exception occurred - {type(e).__name__}")
            
            try:
                # Try clicking with XPATH as fallback
                driver.find_element(By.XPATH, '//*[@id="deny"]').click()
            except Exception as inner_e:
                logger.debug(f"Debug : Failed to click deny button - {type(inner_e).__name__}")
                t_prices.append('//')
                print('*/')
                driver.implicitly_wait(500)
                return 1
        else:
            # Default action if deny button is clicked successfully
            btn.click()
    except TimeoutException:
        logger.debug("Debug : DenyButtonNotFound [In_Time]")
        t_prices.append('//')
        print('/*')
        return 1
    finally:
        # Switch back to the main content in all cases
        driver.switch_to.default_content()

def save_progress_digi(model, price, results_digi ,results_file_digi="results_digi.json", progress_file_digi="progress_digi.json"):
    results_digi[model] = {
        "model": model,
        "price": price
    }
    with open(results_file_digi, "w", encoding="utf-8") as f:
        json.dump(results_digi, f, ensure_ascii=False, indent=2)

    scraped_models = list(results_digi.keys())
    with open(progress_file_digi, "w", encoding="utf-8") as f:
        json.dump(scraped_models, f, ensure_ascii=False, indent=2)

def save_progress_techno(model, price, results_techno ,results_file_techno="results_techno.json", progress_file_techno="progress_techno.json"):
    results_techno[model] = {
        "model": model,
        "price": price
    }
    with open(results_file_techno, "w", encoding="utf-8") as f:
        json.dump(results_techno, f, ensure_ascii=False, indent=2)

    scraped_models = list(results_techno.keys())
    with open(progress_file_techno, "w", encoding="utf-8") as f:
        json.dump(scraped_models, f, ensure_ascii=False, indent=2)


def digi_scrape(driver, digi_scraped, results_digi, results_file_digi, progress_file_digi, resume=False):
    print("Digikala scraping started...")
    for model , url in digi_urls.items():
        if resume:
            if model in digi_scraped:
                d = results_digi[model]
                d_prices.append(d["price"])
                print(f"{GREEN}[✓]{RESET} {model} skipping...")
                continue


        out_off_stock = True
        rang = False

        if url == r"Not_Found": 
            out_off_stock = True
            d_prices.append("**")
            print('**')
            save_progress_digi(model, "**", results_digi, results_file_digi, progress_file_digi)
            continue

        if not wait_for_connection(max_retries=10, retry_delay=10):
            print("Could not establish connection. Exiting program.")
            return 1
        else:
            if not driver.service.process:
                print("Driver instance is invalid.")
                return
            driver.get(url)

        try:
            product_title = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "[data-testid='pdp-title']"))
                )

            try:
                driver.find_element(By.XPATH , '//*[@id="__next"]/div[1]/div[3]/div[3]/div[2]/div[2]/div[2]/div[1]/div/h1/span')
            except NoSuchElementException:
                out_off_stock = False
            else:
                print(f"{model} **")
                d_prices.append('**')
                save_progress_digi(model, "**", results_digi, results_file_digi, progress_file_digi)
                continue
            # checking for the colors available
            try:
                driver.find_element(By.CSS_SELECTOR, "[style='background: rgb(33, 33, 33);']").click()
            except NoSuchElementException:
                try:
                    driver.find_element(By.CSS_SELECTOR, "[style='background: rgb(0, 33, 113);']").click()
                except NoSuchElementException:
                    pass
                else:
                    rang = "Dark Blue"
            else:
                rang = "Black"

            if rang:
                print(model , rang, end=" ")
            else:
                print(model , end=" ")

            try:
                price_no_discount = driver.find_element(By.CSS_SELECTOR , '[data-testid="price-no-discount"]')
            except NoSuchElementException:
                try:
                    final_price_list = driver.find_elements(By.CSS_SELECTOR , '[data-testid="price-final"]')
                    price = final_price_list[1]
                except NoSuchElementException:
                    d_prices.append("//")
                    print('//')
                    save_progress_digi(model, "//", results_digi, results_file_digi, progress_file_digi)
            else:
                if "line-trough" in price_no_discount.get_attribute('class'):
                    final_price_list = driver.find_elements(By.CSS_SELECTOR , '[data-testid="price-final"]')
                    price = final_price_list[1]
                else:
                    price = price_no_discount


            if out_off_stock == False:
                if isinstance(price , str):
                    d_prices.append(price)
                    print(price)
                    save_progress_digi(model, price, results_digi, results_file_digi, progress_file_digi)
                else:
                    final = digits.convert_to_en(price.text)
                    d_prices.append(final)
                    print(final)
                    save_progress_digi(model, final, results_digi, results_file_digi, progress_file_digi)
        except TimeoutException:
            print(f"{RED}[!]{RESET} Failed to find the title for {model} within the given time.")
            d_prices.append('//')

        continue
        # d_pbar.update(1)
    print(f"Digikala scraping Finished{GREEN}[✓]{RESET}")
    driver.quit()


percent = 100 / len(techno_urls)

# loading the page
def techno_scrape(driver, techno_scraped, results_techno, results_file_techno, progress_file_techno, resume=False):
    print("Techno Life scraping started...")
    for model , url in techno_urls.items():
        if resume:
            if model in techno_scraped:
                t = results_techno[model]
                t_prices.append(t["price"])
                print(f"{GREEN}[✓]{RESET} {model} skipping...")
                continue


        print(model , end="---")

        if url == r"Not_Found": 
            out_off_stock = True
            t_prices.append("**")
            print('**')
            save_progress_techno(model, "**", results_techno, results_file_techno, progress_file_techno)
            continue

        if not wait_for_connection(max_retries=10, retry_delay=10):
            print(f"{RED}[!]{RESET}Could not establish connection. Exiting program.")
            return 1
        else:
            if not driver.service.process:
                print(f"{RED}[!]{RESET} Driver instance is invalid.")
                return
            driver.get(url)

        try:
            product_title = WebDriverWait(driver, 20).until(
                    EC.presence_of_element_located((By.ID, "pdp_name"))
                )

            try:
                out_off_stock = driver.find_element(By.XPATH , '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[2]/div/div/div/div/div/p[contains (text() , "ناموجود")]')
            except NoSuchElementException:
                pass
            else:
                t_prices.append("**")
                print('**')
                save_progress_techno(model, "**", results_techno, results_file_techno, progress_file_techno)
                continue

            out_off_stock = False
            price = "//"


            # checking for the colors available
            try:
                black_btn = driver.find_element(By.CSS_SELECTOR, "[style='background-color:#1a1a1a'")
            except NoSuchElementException:
                try:
                    dark_blue_btn = driver.find_element(By.CSS_SELECTOR, "[style='background-color:#00009c']")
                except NoSuchElementException:
                    pass
                else:
                    try:
                        dark_blue_btn.click()
                    except ElementClickInterceptedException:
                        if deny(dark_blue_btn, driver) == 1:
                            continue
                    finally:
                        rang = "DarkBlue"
            else:
                try:
                    black_btn.click()
                except ElementClickInterceptedException:
                    if deny(black_btn, driver) == 1:
                        continue
                finally:
                    rang = "Black"


            # finding the price and scraping it
            for x in xpath_for_price_techno:
                try:
                    price = driver.find_element(By.XPATH , xpath_for_price_techno[x])
                except NoSuchElementException:
                    pass
                else:
                    break

            if out_off_stock == False:
                if isinstance(price, str):
                    t_prices.append(price)
                    print(price)
                    save_progress_techno(model, price , results_techno, results_file_techno, progress_file_techno)
                else:
                    t_prices.append(price.text)
                    print(price.text)
                    save_progress_techno(model, price.text, results_techno, results_file_techno, progress_file_techno)
        except TimeoutException:
                print(f"[!] Failed to find the title for {model} within the given time.")
                t_prices.append('//')
        continue
        # t_pbar.update(1)
    print(f"Techno Life scraping Finished{GREEN}[✓]{RESET}")
    driver.quit()


def create_document():
    # creating the document and the row
    document = Document()
    table = document.add_table(rows=1, cols=3)

    # giving the style values

    style = document.styles['Normal']
    table.style = 'Table Grid'
    style.font.name = "Calibri" # type: ignore
    style.font.size = Pt(20) # type: ignore

    # The array of Phone model names , digikala prices , technolife prices


    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('phone').bold = True
    hdr_cells[1].text = 'Digikala'
    hdr_cells[2].text = 'Technolife'

    for i in range(urls_len):
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(phone_models[i]).bold = True
        row_cells[1].text = d_prices[i]
        row_cells[2].text = t_prices[i]

    today_date = str(JalaliDate.today())
    now_time = datetime.now().strftime("%H_%M")
    file_name = f"{today_date[5:]} {now_time}"
    path = today_date[:-3]


    if not os.path.exists(path):
        os.makedirs(path)

    doc_file = os.path.join(path, f"{file_name}.docx")
    document.save(doc_file)

    # Convert the document to PDF
    pdf_file = os.path.join(path, f'{file_name}.pdf')
    convert(doc_file, pdf_file)
    # output_file = os.path.join(,{})
    output_file = pdf_file

    os.remove(doc_file)

    if output_file is None:
        raise ValueError("Error: PDF conversion failed! convert() returned None.")


    print(f"Generated PDF path: {output_file}, Exists: {os.path.exists(output_file)}")
    return output_file

def RunTest(t, d):
    t.append("123")
    d.append("123")

    try:
        prices_pdf = create_document()
    except TimeoutError as e:
        print(f"Creating the Document failed with TimeOutError : {e}")
        return 1
    except Exception as e:
        print(f"Creating the Document failed with this error : {e}")
        return 1

    try:
        send_to_Eitaa(prices_pdf)
    except TimeoutError as e:
        print(f"sending file to Eitaa failed with TimeOutError : {e}")
        return 1
    except Exception as e:
        print(f"sending file to Eitaa failed with this error : {e}")
        return 1


def main(resume=False):

    if Test_Mode:
        try:
            RunTest(t_prices, d_prices)
        except TimeoutError as e: 
            print(f"{RED}[!]{RESET} RunTest failed with this TimeOutError : {e}")
        except Exception as e:
            print(f"{RED}[!]{RESET} RunTest failed with this error : {e}")
        finally:
            return 0

    if not resume:
        for file in ["progress_digi.json", "results_digi.json", "progress_techno.json", "results_techno.json"]:
            try:
                os.remove(file)
                print(f"{GREEN}[✓]{RESET}{file} removed")
            except FileNotFoundError:
                print(f"{RED}[!]{RESET}{file} Not Found")
                pass
            except Exception as e:
                print(f"{RED}[!]{RESET}Failed to remove {file} due to this error: {e}")

    progress_file_digi = "progress_digi.json"
    results_file_digi = "results_digi.json"


    if os.path.exists(progress_file_digi):
        with open(progress_file_digi, "r") as f:
            digi_scraped = set(json.load(f))
    else:
        digi_scraped = set()

    if not os.path.exists(results_file_digi):
        with open(results_file_digi, "w", encoding="utf-8") as f:
            json.dump({}, f, ensure_ascii=False, indent=2)

    with open(results_file_digi, "r", encoding="utf-8") as f:
        results_digi = json.load(f)


    try:
        driver = driver_setup()
        digi_start = time.time()
        result = digi_scrape(driver, digi_scraped, results_digi, results_file_digi, progress_file_digi, resume)
        digi_end = time.time()
        print((digi_end - digi_start) / 60)
        if result == 1:
            raise SystemExit(f"{RED}[!]{RESET} Critical Error: digi scraping failed. Exiting the app...")
    except TimeoutError as e:
        print(f"{RED}[!]{RESET} Digi scraping failed with Timeout Error")
        return 1
    except Exception as e:
        print(f"{RED}[!]{RESET} Digi scraping failed with this error : {e}")
        traceback.print_exc()
        return 1
    finally:
        driver.quit()

    progress_file_techno = "progress_techno.json"
    results_file_techno = "results_techno.json"

    if os.path.exists(progress_file_techno):
        with open(progress_file_techno, "r") as f:
            techno_scraped = set(json.load(f))
    else:
        techno_scraped = set()

    if not os.path.exists(results_file_techno):
        with open(results_file_techno, "w", encoding="utf-8") as f:
            json.dump({}, f, ensure_ascii=False, indent=2)

    with open(results_file_techno, "r", encoding="utf-8") as f:
        results_techno = json.load(f)


    try:
        driver = driver_setup()
        techno_start = time.time()
        result = techno_scrape(driver, techno_scraped, results_techno, results_file_techno ,progress_file_techno, resume)
        techno_end = time.time()
        print((techno_end - techno_start) / 60)
        if result == 1:
            raise SystemExit(f"{RED}[!]{RESET} Critical Error: techno scraping failed. Exiting the app...")
    except TimeoutError as e:
        print(f"{RED}[!]{RESET} techno scraping failed with Timeout Error")
        return 1
    except Exception as e:
        print(f"{RED}[!]{RESET} techno scraping failed with this error : {e}")
        traceback.print_exc()
        return 1
    finally:
        driver.quit()

    try:
        prices_pdf = create_document()
    except TimeoutError as e:
        print(f"{RED}[!]{RESET} Creating the Document failed with Timeout Error")
        return 1
    except Exception as e:
        print(f"{RED}[!]{RESET} Creating the Document failed with this error : {e}")
        traceback.print_exc()
        return 1

    try:
        send_to_Eitaa(prices_pdf)
    except TimeoutError as e:
        print(f"{RED}[!]{RESET} sending file to Eitaa failed with this error : {e}")
        return 1
    except Exception as e:
        print(f"{RED}[!]{RESET} sending file to Eitaa failed with this error : {e}")
        return 1
    finally:
        driver.quit()

    os.remove("progress_digi.json")
    os.remove("results_digi.json")
    os.remove("progress_techno.json")
    os.remove("results_techno.json")




if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Script to scrape and create documents.")

    # --resume argument
    parser.add_argument("--resume", action="store_true", help="Enable resume mode")

    args = parser.parse_args()
    main(resume=args.resume)